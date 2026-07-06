"""Extract normalized text and metadata from supported document formats."""

from __future__ import annotations

import hashlib
from dataclasses import dataclass
from pathlib import Path


SUPPORTED_TYPES = {".txt": "txt", ".pdf": "pdf", ".docx": "docx"}


class DocumentLoadError(ValueError):
    """Raised when an uploaded document cannot be safely converted to text."""


@dataclass(frozen=True)
class LoadedDocument:
    text: str
    source_type: str
    original_filename: str
    source_hash: str


def _extract_txt(path: Path) -> str:
    data = path.read_bytes()
    if b"\x00" in data:
        raise DocumentLoadError("TXT file appears to contain binary data.")
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise DocumentLoadError("TXT file must use UTF-8 encoding.") from error


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        if reader.is_encrypted:
            raise DocumentLoadError("Encrypted PDF files are not supported.")
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    except DocumentLoadError:
        raise
    except Exception as error:
        raise DocumentLoadError("PDF file is unreadable or invalid.") from error


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document

        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as error:
        raise DocumentLoadError("DOCX file is unreadable or invalid.") from error


def load_document(path: Path, *, original_filename: str | None = None) -> LoadedDocument:
    """Load one supported file without executing or trusting its contents."""
    path = path.resolve()
    if not path.is_file():
        raise DocumentLoadError(f"Document file does not exist: {path}")

    extension = path.suffix.lower()
    source_type = SUPPORTED_TYPES.get(extension)
    if source_type is None:
        raise DocumentLoadError("Unsupported file type. Use TXT, PDF, or DOCX.")

    extractors = {
        "txt": _extract_txt,
        "pdf": _extract_pdf,
        "docx": _extract_docx,
    }
    text = extractors[source_type](path)
    normalized = "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").split("\n"))
    if not normalized.strip():
        detail = " The PDF may be scanned and require OCR." if source_type == "pdf" else ""
        raise DocumentLoadError(f"Document does not contain extractable text.{detail}")

    display_name = Path(original_filename or path.name).name
    return LoadedDocument(
        text=normalized,
        source_type=source_type,
        original_filename=display_name,
        source_hash=hashlib.sha256(path.read_bytes()).hexdigest(),
    )
