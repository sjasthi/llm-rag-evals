"""Unit tests for deterministic, provider-free RAG helper behavior."""

from __future__ import annotations

import sys
import tempfile
import unittest
from dataclasses import replace
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch


PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "rag"))

from answer import answer_question  # noqa: E402
from document_loader import DocumentLoadError, load_document  # noqa: E402
from ingest import build_loaded_document_chunks, chunk_text, stable_chroma_id  # noqa: E402
from llm import REFUSAL_MESSAGE, SYSTEM_INSTRUCTION, build_grounded_prompt, generate_with_gemini  # noqa: E402
from query import SearchResult, lexical_score  # noqa: E402
from settings import load_settings  # noqa: E402


class ChunkTextTests(unittest.TestCase):
    def test_empty_text_produces_no_chunks(self) -> None:
        self.assertEqual([], chunk_text("  \n\t", chunk_size=20, overlap=5))

    def test_overlap_is_preserved(self) -> None:
        self.assertEqual(
            ["abcdefghij", "hijklmnopq", "opqrstuvwx", "vwxyz"],
            chunk_text("abcdefghijklmnopqrstuvwxyz", chunk_size=10, overlap=3),
        )

    def test_invalid_overlap_is_rejected(self) -> None:
        with self.assertRaises(ValueError):
            chunk_text("text", chunk_size=10, overlap=10)


class StableIdTests(unittest.TestCase):
    def test_same_path_and_index_produce_same_id(self) -> None:
        first = stable_chroma_id("data/example.txt", 3)
        second = stable_chroma_id("data/example.txt", 3)
        self.assertEqual(first, second)

    def test_chunk_index_changes_id(self) -> None:
        self.assertNotEqual(
            stable_chroma_id("data/example.txt", 2),
            stable_chroma_id("data/example.txt", 3),
        )


class DocumentLoaderTests(unittest.TestCase):
    def test_utf8_txt_is_loaded_with_safe_original_name(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "stored.txt"
            path.write_text("Metro State registration details", encoding="utf-8-sig")
            loaded = load_document(path, original_filename="../Registration.txt")

        self.assertEqual("txt", loaded.source_type)
        self.assertEqual("Registration.txt", loaded.original_filename)
        self.assertIn("registration details", loaded.text)

    def test_binary_txt_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "bad.txt"
            path.write_bytes(b"text\x00binary")
            with self.assertRaisesRegex(DocumentLoadError, "binary data"):
                load_document(path)

    def test_empty_pdf_is_rejected_as_non_extractable(self) -> None:
        fake_reader = SimpleNamespace(is_encrypted=False, pages=[SimpleNamespace(extract_text=lambda: "")])
        fake_module = SimpleNamespace(PdfReader=lambda _path: fake_reader)
        with tempfile.TemporaryDirectory() as directory, patch.dict(sys.modules, {"pypdf": fake_module}):
            path = Path(directory) / "scan.pdf"
            path.write_bytes(b"%PDF fixture")
            with self.assertRaisesRegex(DocumentLoadError, "require OCR"):
                load_document(path)

    def test_pdf_text_is_extracted(self) -> None:
        fake_reader = SimpleNamespace(
            is_encrypted=False,
            pages=[SimpleNamespace(extract_text=lambda: "Financial aid deadline")],
        )
        fake_module = SimpleNamespace(PdfReader=lambda _path: fake_reader)
        with tempfile.TemporaryDirectory() as directory, patch.dict(sys.modules, {"pypdf": fake_module}):
            path = Path(directory) / "aid.pdf"
            path.write_bytes(b"%PDF fixture")
            loaded = load_document(path)
        self.assertEqual("pdf", loaded.source_type)
        self.assertEqual("Financial aid deadline", loaded.text)

    def test_docx_paragraph_and_table_are_extracted(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "support.docx"
            document = Document()
            document.add_paragraph("Student support services")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Office"
            table.cell(0, 1).text = "Library"
            document.save(path)
            loaded = load_document(path)

        self.assertIn("Student support services", loaded.text)
        self.assertIn("Office\tLibrary", loaded.text)

    def test_loaded_metadata_is_preserved_in_chunks(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "policy.txt"
            path.write_text("A policy statement long enough for two chunks.", encoding="utf-8")
            loaded = load_document(path, original_filename="Policy.txt")
            chunks = build_loaded_document_chunks(
                loaded,
                source_path="storage/uploads/id.txt",
                category="policies",
                chunk_size=30,
                overlap=5,
            )

        self.assertGreater(len(chunks), 1)
        self.assertTrue(all(chunk.source_type == "txt" for chunk in chunks))
        self.assertTrue(all(chunk.original_filename == "Policy.txt" for chunk in chunks))


class LexicalScoreTests(unittest.TestCase):
    def test_exact_answer_phrase_scores_above_partial_match(self) -> None:
        question = "When does Fall 2026 registration begin?"
        exact = "Fall 2026 registration begins in eServices on March 23."
        partial = "Fall visiting student registration opens in June 2026."
        self.assertGreater(lexical_score(question, exact), lexical_score(question, partial))


def sample_context() -> SearchResult:
    return SearchResult(
        rank=1,
        source_path="data/metrostate_documents/academic_calendar/fall_2026.txt",
        category="academic_calendar",
        chunk_index=0,
        text="Fall 2026 registration begins in eServices Monday, March 23, 2026.",
        distance=0.5,
        keyword_score=10,
        document_id=2,
        chunk_id=10,
    )


class GroundedAnswerTests(unittest.TestCase):
    def test_prompt_contains_question_and_source_metadata(self) -> None:
        prompt = build_grounded_prompt("When does registration begin?", [sample_context()])
        self.assertIn("When does registration begin?", prompt)
        self.assertIn("fall_2026.txt", prompt)
        self.assertIn("March 23, 2026", prompt)

    def test_system_instruction_requires_exact_refusal(self) -> None:
        self.assertIn(REFUSAL_MESSAGE, SYSTEM_INSTRUCTION)
        self.assertIn("only the supplied source excerpts", SYSTEM_INSTRUCTION)

    def test_answer_workflow_accepts_injected_generator(self) -> None:
        def fake_generator(prompt: str, _settings: object) -> str:
            self.assertIn("March 23, 2026", prompt)
            return "Registration begins March 23, 2026."

        with patch("answer.semantic_search", return_value=[sample_context()]):
            result = answer_question(
                "When does Fall 2026 registration begin?",
                top_k=1,
                save=False,
                generator=fake_generator,
            )

        self.assertEqual("Registration begins March 23, 2026.", result.answer)
        self.assertEqual(1, len(result.sources))
        self.assertIsNone(result.response_id)

    def test_gemini_requires_api_key(self) -> None:
        settings = replace(load_settings(), llm_api_key="")
        with self.assertRaisesRegex(RuntimeError, "API key is not configured"):
            generate_with_gemini("prompt", settings)


if __name__ == "__main__":
    unittest.main()
